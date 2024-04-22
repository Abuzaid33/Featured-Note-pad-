import tkinter as tk
from tkinter import ttk, messagebox, filedialog, simpledialog
from docx import Document
import speech_recognition as sr
from PIL import Image, ImageTk
import pytesseract
import os
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from tkinter import colorchooser
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from fpdf import FPDF
import vosk
import pyaudio
import json
import threading
###############################
# Section: Start page
###############################

class StartPage:
    def __init__(self, master, on_start_button_click):
        self.master = master
        self.start_window = tk.Toplevel(master)
        self.start_window.title("Welcome to NoteCraft")
        bg_image = Image.open(os.path.abspath("images/pic.png")).resize((master.winfo_screenwidth(), master.winfo_screenheight()))
        self.bg_photo = ImageTk.PhotoImage(bg_image)
        background_label = tk.Label(self.start_window, image=self.bg_photo)
        background_label.image = self.bg_photo 
        background_label.place(relwidth=1, relheight=1)
        self.start_window.geometry(f"{master.winfo_screenwidth()}x{master.winfo_screenheight()}+0+0")
        self.start_window.configure(bg='#d7bde2')  
        style = ttk.Style()
        style.configure("Start.TButton",
                        font=('Helvetica', 16, 'bold'),
                        foreground='black',
                        background='black',  
                        borderwidth=0,
                        focuscolor='black',
                        lightcolor='black',
                        darkcolor='black',
                        padding=(15, 10))
        start_label = ttk.Label(self.start_window, text="Welcome to NoteCraft!", font=('Arial', 24), foreground='white', background='#d7bde2')  # Light purple background
        start_label.place(relx=0.5, rely=0.4, anchor="center")
        icon_path = "images/download.png"
        self.start_window.iconbitmap(icon_path)
        start_button = ttk.Button(self.start_window, text="Start", command=on_start_button_click, style="Start.TButton")
        start_button.place(relx=0.5, rely=0.6, anchor="center")
          
class NoteCraftApp:
    def __init__(self, root,model_path):
        self.root = root
        self.root.title("NoteCraft")
        self.root.state('zoomed')
        self.root.withdraw() 
        self.start_page = StartPage(self.root, self.initialize_main_app)
        bg_image = Image.open(os.path.abspath("images/picture.png")).resize((root.winfo_screenwidth(), root.winfo_screenheight()))
        self.bg_photo = ImageTk.PhotoImage(bg_image)
        background_label = tk.Label(root, image=self.bg_photo)
        background_label.place(relwidth=1, relheight=1)
        button_frame = ttk.Frame(root, style='TFrame', padding=(10, 5, 10, 5), relief='flat', borderwidth=2)
        button_frame.pack(pady=10, anchor='center')
        icon_path = ("images/download.png")
        self.root.iconbitmap("images/download.png")
        self.model_path = model_path
        self.all_notes_window = None
        self.desktop_path = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
        self.all_files_path = "All Files"
        if not os.path.exists(self.all_files_path):
            os.makedirs(self.all_files_path)

        ###############################
        # Add Scrollbar to Editor
        ###############################
        text_frame = ttk.Frame(root)
        text_frame.pack(padx=1, pady=1, expand=True, fill=tk.BOTH)
        self.note_text = tk.Text(text_frame, wrap="word", font=("Arial", 12),undo=True)
        self.note_text.pack(side=tk.LEFT, expand=True, fill=tk.BOTH)
        scrollbar = ttk.Scrollbar(text_frame, orient=tk.VERTICAL, command=self.note_text.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.note_text.config(yscrollcommand=scrollbar.set)

        ###############################
        # Section: Main  Functions
        ###############################
        
        file_menu_button = ttk.Menubutton(button_frame, text="File")
        edit_menu_button = ttk.Menubutton(button_frame, text="Edit")
        format_menu_button = ttk.Menubutton(button_frame, text="Format")
        file_menu = tk.Menu(file_menu_button, tearoff=0)
        edit_menu = tk.Menu(edit_menu_button, tearoff=0)
        format_menu = tk.Menu(format_menu_button, tearoff=0)
        file_menu_button.configure(menu=file_menu)
        edit_menu_button.configure(menu=edit_menu)
        format_menu_button.configure(menu=format_menu)
        file_menu_button.grid(row=0, column=0, padx=5)
        edit_menu_button.grid(row=0, column=1, padx=5)
        format_menu_button.grid(row=0, column=2, padx=5)
        speech_to_text_button = ttk.Menubutton(button_frame, text="Speech to Text", style='TButton')
        speech_to_text_button.grid(row=0, column=3, padx=5)
        speech_to_text_menu = tk.Menu(speech_to_text_button, tearoff=0)
        speech_to_text_button.configure(menu=speech_to_text_menu)
        speech_to_text_menu.add_command(label="Speech to Text", command=self.start_speech_to_text_thread)
        image_to_text_button = ttk.Button(button_frame, text="Image to Text", command=self.image_to_text, style='TButton')
        all_notes_button = ttk.Button(button_frame, text="All Notes", command=self.show_all_notes, style='TButton')
        image_to_text_button.grid(row=0, column=4, padx=5)
        all_notes_button.grid(row=0, column=5, padx=5)
        ###############################
        # Section: File Functions
        ###############################
        
        file_menu.add_command(label="New Note", command=self.new_note)
        file_menu.add_command(label="Open Note", command=self.open_note)
        file_menu.add_command(label="Save Note", command=self.save_note)
        file_menu.add_command(label="Save As", command=self.save_as)
        file_menu.add_command(label="Save as PDF", command=self.save_as_pdf)
        file_menu.add_command(label="Save as DOCX", command=self.save_as_docx)
        file_menu.add_command(label="PDF (Custom saving)", command=self.save_as_pdf_custom)
        file_menu.add_command(label="DOCX (Custom saving)", command=self.save_as_docx_custom)
        file_menu.add_command(label="Exit", command=self.exit_app)
        
        ###############################
        # Section: Edit Functions
        ###############################
        
        edit_menu.add_command(label="Cut", command=self.cut_text)
        edit_menu.add_command(label="Copy", command=self.copy_text)
        edit_menu.add_command(label="Paste", command=self.paste_text)
        edit_menu.add_command(label="Undo", command=self.note_text.edit_undo)
        edit_menu.add_command(label="Redo", command=self.note_text.edit_redo)
      
        ###############################
        # Section: Format Functions
        ###############################

        format_menu.add_command(label="Bold", command=self.bold_text)
        format_menu.add_command(label="Italic", command=self.italic_text)
        format_menu.add_command(label="Underline", command=self.underline_text)
        format_menu.add_command(label="Font Size", command=self.set_font_size)
        format_menu.add_command(label="Font Color", command=self.change_font_color)

        ###############################
        # Section: All Notes Functions
        ###############################

        self.all_notes_frame = ttk.Frame(root, style='TFrame', padding=(10, 5, 10, 5), relief='flat', borderwidth=2)
        self.all_notes_frame.pack(pady=10)
        self.all_notes_window = None

    ###############################
    # Section: File Functions
    ###############################

    def new_note(self):
        self.note_text.delete(1.0, tk.END)
        messagebox.showinfo("New Note", "New note created successfully!")

    def open_note(self):
        file_path = filedialog.askopenfilename(defaultextension=".txt", filetypes=[("Text files", "*.txt")])
        if file_path:
            with open(file_path, 'r') as file:
                content_tags = []
                for line in file:
                    parts = line.strip().split("::")
                    if len(parts) > 1:
                        content_tags.append((parts[0], parts[1]))
                    else:
                        content_tags.append((parts[0], None))
            self.note_text.delete("1.0", tk.END)
            for text, tag in content_tags:
                self.note_text.insert(tk.END, text)
                if tag:
                    self.note_text.tag_add(tag, "end-1c", "end")
            messagebox.showinfo("Open Note", f"Note opened from {file_path}")

    def save_note(self, event=None):
        file_name = simpledialog.askstring("Save Note", "Enter a name for the note:")
        if file_name:
            file_path_desktop = os.path.join(self.desktop_path, file_name + ".txt")
            file_path_all_files = os.path.join(self.all_files_path, file_name + ".txt")
            if os.path.exists(file_path_desktop) or os.path.exists(file_path_all_files):
                response = messagebox.askyesno("Overwrite File", f"The file '{file_name}.txt' already exists. Do you want to overwrite it?")
                if not response:
                    return
            with open(file_path_desktop, 'w') as file_desktop:
                for index, char in enumerate(self.note_text.get("1.0", tk.END)):
                    tag_names = self.note_text.tag_names(f"{index + 1}.0")
                    if tag_names:
                        formatted_tags = "::".join(tag_names)
                        file_desktop.write(f"{char}::{formatted_tags}\n")
                    else:
                        file_desktop.write(char)
            with open(file_path_all_files, 'w') as file_all_files:
                for index, char in enumerate(self.note_text.get("1.0", tk.END)):
                    tag_names = self.note_text.tag_names(f"{index + 1}.0")
                    if tag_names:
                        formatted_tags = "::".join(tag_names)
                        file_all_files.write(f"{char}::{formatted_tags}\n")
                    else:
                        file_all_files.write(char)
            messagebox.showinfo("Save Note", f"Note saved to {file_path_desktop} and {file_path_all_files}")
            self.current_file_path = file_path_desktop
            if not hasattr(self, 'current_file_name'):
                self.populate_all_notes()
        root.bind('<Control-s>', self.save_note)

    def save_as(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text files", "*.txt")])
        if file_path:
            with open(file_path, 'w') as file:
                content = self.note_text.get(1.0, tk.END)
                file.write(content)
            messagebox.showinfo("Save As", f"Note saved as {file_path}")   
    
    def save_as_pdf(self):
        file_name = simpledialog.askstring("Save as PDF", "Enter a name for the PDF:")
        if file_name:
            file_path_desktop = os.path.join(self.desktop_path, file_name + ".pdf")
            file_path_all_files = os.path.join(self.all_files_path, file_name + ".pdf")
            if os.path.exists(file_path_desktop) or os.path.exists(file_path_all_files):
                response = messagebox.askyesno("Overwrite File", f"The file '{file_name}.pdf' already exists. Do you want to overwrite it?")
                if not response:
                    return 
            content = self.note_text.get(1.0, tk.END)
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 10, content)
            pdf.output(file_path_desktop)
            pdf.output(file_path_all_files)
            messagebox.showinfo("Save as PDF", f"Note saved as PDF to {file_path_desktop} and {file_path_all_files}")
            self.populate_all_notes()

    def save_as_docx(self):
        file_name = simpledialog.askstring("Save as DOCX", "Enter a name for the DOCX:")
        if file_name:
            file_path_desktop = os.path.join(self.desktop_path, file_name + ".docx")
            file_path_all_files = os.path.join(self.all_files_path, file_name + ".docx")
            if os.path.exists(file_path_desktop) or os.path.exists(file_path_all_files):
                response = messagebox.askyesno("Overwrite File", f"The file '{file_name}.docx' already exists. Do you want to overwrite it?")
                if not response:
                    return 
            content = self.note_text.get(1.0, tk.END)
            doc = Document()
            self.add_paragraph_with_formatting(doc, content)
            doc.save(file_path_desktop)
            doc.save(file_path_all_files)
            messagebox.showinfo("Save as DOCX", f"Note saved as DOCX to {file_path_desktop} and {file_path_all_files}")
            self.populate_all_notes()

    def add_paragraph_with_formatting(self, doc, text, bold=False, italic=False, underline=False, color=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        font = run.font
        font.size = Pt(12)  
        font.bold = bold   
        font.italic = italic  
        font.underline = underline  
        if color is not None:
            if isinstance(color, tuple) and len(color) == 3:
                rgb_color = RGBColor(*color)
            elif isinstance(color, str):
                rgb_color = WD_COLOR_INDEX[color.upper()]
            else:
                raise ValueError("Invalid color format. Provide an RGB tuple or a color name string.")
            font.color.rgb = rgb_color 
        return p

    def save_as_pdf_custom(self):
        file_name = simpledialog.askstring("Save as PDF", "Enter a name for the PDF:")
        if file_name:
            file_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")], initialfile=file_name)
            if file_path:
                pdf_canvas = canvas.Canvas(file_path, pagesize=letter)
                content = self.note_text.get(1.0, tk.END)
                pdf_canvas.drawString(100, 800, content)
                pdf_canvas.save()
                messagebox.showinfo("Save as PDF", f"Note saved as PDF to {file_path}")
                self.populate_all_notes()

    def save_as_docx_custom(self):
        file_name = simpledialog.askstring("Save as DOCX", "Enter a name for the DOCX:")
        if file_name:
            file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")], initialfile=file_name)
            if file_path:
                doc = Document()
                content = self.note_text.get(1.0, tk.END)
                doc.add_paragraph(content)
                doc.save(file_path)
                messagebox.showinfo("Save as DOCX", f"Note saved as DOCX to {file_path}")
                self.populate_all_notes()

    def exit_app(self):
        self.root.destroy()

    ###############################
    # Section: Edit Functions
    ###############################

    def edit_menu(self):
        if self.cut_button.winfo_ismapped():
            self.hide_edit_buttons()
        else:
            self.show_edit_buttons()

    def cut_text(self):
        selected_text = self.note_text.get("sel.first", "sel.last")
        if selected_text:
            self.root.clipboard_clear()  
            self.root.clipboard_append(selected_text)  
            self.note_text.delete("sel.first", "sel.last")

    def copy_text(self):
        if self.note_text.tag_ranges(tk.SEL):
            selected_text = self.note_text.get(tk.SEL_FIRST, tk.SEL_LAST)
            self.root.clipboard_clear()
            self.root.clipboard_append(selected_text)

    def paste_text(self):
        try:
            clipboard_content = self.root.clipboard_get()
            if clipboard_content:
                self.note_text.insert(tk.INSERT, clipboard_content)
        except tk.TclError:
            pass
 
    ###############################
    # Section: Format Functions
    ###############################

    def bold_text(self):
        current_tags = self.note_text.tag_names(tk.SEL_FIRST)
        if "bold" not in current_tags:
            self.note_text.tag_add("bold", tk.SEL_FIRST, tk.SEL_LAST)
            self.note_text.tag_configure("bold", font=('Arial', 16, 'bold'))
        else:
            self.note_text.tag_remove("bold", tk.SEL_FIRST, tk.SEL_LAST)

    def italic_text(self):
        current_tags = self.note_text.tag_names(tk.SEL_FIRST)
        if "italic" not in current_tags:
            self.note_text.tag_add("italic", tk.SEL_FIRST, tk.SEL_LAST)
            self.note_text.tag_configure("italic", font=('Arial', 16, 'italic'))
        else:
            self.note_text.tag_remove("italic", tk.SEL_FIRST, tk.SEL_LAST)

    def underline_text(self):
        current_tags = self.note_text.tag_names(tk.SEL_FIRST)
        if "underline" not in current_tags:
            self.note_text.tag_add("underline", tk.SEL_FIRST, tk.SEL_LAST)
            self.note_text.tag_configure("underline", underline=True)
        else:
            self.note_text.tag_remove("underline", tk.SEL_FIRST, tk.SEL_LAST)

    def set_font_size(self):
        size = simpledialog.askinteger("Font Size", "Enter font size:")
        if size:
            current_tags = self.note_text.tag_names(tk.SEL_FIRST)
            if "font_size" in current_tags:
                self.note_text.tag_remove("font_size", tk.SEL_FIRST, tk.SEL_LAST)
            self.note_text.tag_add("font_size", tk.SEL_FIRST, tk.SEL_LAST)
            self.note_text.tag_configure("font_size", font=('Arial', size))
            
    def change_font_color(self):
        color = colorchooser.askcolor(title="Choose Font Color")[1]
        if color:
            self.note_text.tag_add("font_color", tk.SEL_FIRST, tk.SEL_LAST)
            self.note_text.tag_configure("font_color", foreground=color)
    
    ###############################
    # Section: Speech to Text offline Functions
    ###############################
            
    def start_speech_to_text_thread(self):
        threading.Thread(target=self.speech_to_text).start()

    def speech_to_text(self):
        recognizer = vosk.KaldiRecognizer(vosk.Model(self.model_path), 16000)
        audio = pyaudio.PyAudio()
        stream = audio.open(format=pyaudio.paInt16, channels=1, rate=16000, input=True, frames_per_buffer=1024)
        print("Listening...")
        while True:
            try:
                data = stream.read(1024)
                if len(data) == 0:
                    break
                if recognizer.AcceptWaveform(data):
                    result = recognizer.Result()
                    result_dict = json.loads(result)
                    if 'text' in result_dict:
                        text = result_dict['text']
                        self.root.after(0, lambda: self.note_text.insert(tk.END, text))
            except KeyboardInterrupt:
                break
            except Exception as e:
                print(f"Error: {e}")
        print("Stopped listening.")
        stream.stop_stream()
        stream.close()
        audio.terminate()
    
    ###############################
    # Section: Image to Text Functions
    ###############################

    def image_to_text(self):
        image_path = filedialog.askopenfilename(filetypes=[("Image files", "*.png;*.jpg;*.jpeg;*.gif")])
        if image_path:
            try:
                image = Image.open(image_path)
                text = pytesseract.image_to_string(image)
                self.note_text.insert(tk.END, text + "\n")
            except Exception as e:
                messagebox.showerror("Error", f"An error occurred: {e}")

    ###############################
    # Section: All Notes Functions
    ###############################

    def show_all_notes(self):
        if not self.all_notes_window or not self.all_notes_window.winfo_exists():
            self.all_notes_window = tk.Toplevel(self.root)
            self.all_notes_window.title("All Notes")
            self.populate_all_notes()
        else:
            self.all_notes_window.lift()

    def populate_all_notes(self):
        all_files_dir = "All Files"
        if not os.path.exists(all_files_dir):
            os.makedirs(all_files_dir)
        files_and_folders = []
        for item in os.listdir(all_files_dir):
            item_path = os.path.join(all_files_dir, item)
            if os.path.isfile(item_path):
                files_and_folders.append({"name": item, "type": "file", "path": item_path})
            elif os.path.isdir(item_path):
                files_and_folders.append({"name": item, "type": "folder", "path": item_path})
        if self.all_notes_window:
            for widget in self.all_notes_window.winfo_children():
                widget.destroy()
        bg_image_all_notes = Image.open(os.path.abspath("images/image.png")).resize((self.root.winfo_screenwidth(), self.root.winfo_screenheight()))
        bg_photo_all_notes = ImageTk.PhotoImage(bg_image_all_notes)
        background_label_all_notes = tk.Label(self.all_notes_window, image=bg_photo_all_notes)
        background_label_all_notes.image = bg_photo_all_notes  
        background_label_all_notes.place(relwidth=1, relheight=1)
        self.all_notes_window.configure(bg='#D3D3D3') 
        self.all_notes_window.geometry(f"{self.root.winfo_screenwidth()}x{self.root.winfo_screenheight()}+0+0")
        self.all_notes_window.update_idletasks() 
        center_x = (self.all_notes_window.winfo_screenwidth() - self.all_notes_window.winfo_width()) // 2
        center_y = (self.all_notes_window.winfo_screenheight() - self.all_notes_window.winfo_height()) // 2
        self.all_notes_window.geometry(f"+{center_x}+{center_y}")
        for note in files_and_folders:
            note_frame = ttk.Frame(self.all_notes_window, style='TFrame', padding=(7, 3, 7, 3), relief='flat', borderwidth=2)
            note_frame.pack(fill=tk.X, padx=7, pady=3)
            note_label_text = f"{note['name']} {'📁' if note['type'] == 'folder' else '📄'}"
            note_label = ttk.Label(note_frame, text=note_label_text, style="Start.TButton")
            note_label.pack(side=tk.TOP, anchor='w', padx=5, pady=5)
            open_button = ttk.Button(note_frame, text="Open", command=lambda n=note: self.open_file_or_folder(n), style="Start.TButton")
            update_button = ttk.Button(note_frame, text="Update", command=lambda n=note: self.update_file(n), style="Start.TButton")
            delete_button = ttk.Button(note_frame, text="Delete", command=lambda n=note: self.delete_file(n), style="Start.TButton")
            rename_button = ttk.Button(note_frame, text="Rename", command=lambda n=note: self.rename_file(n), style="Start.TButton")
            open_button.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=3)
            update_button.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=3)
            delete_button.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=3)
            rename_button.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=3)
        ttk.Frame(self.all_notes_window, height=10, style='TFrame').pack()

    def open_file_or_folder(self, note):
        file_path = note.get("path")
        if file_path:
            try:
                os.startfile(file_path)  
            except Exception as e:
                messagebox.showerror("Error", f"An error occurred while opening the file: {e}")

    def update_file(self, note):
        file_path = note.get("path")
        if file_path:
            try:
                with open(file_path, 'r') as file:
                    current_content = file.read()
                self.note_text.delete(1.0, tk.END)
                self.note_text.insert(tk.END, current_content)
            except Exception as e:
                messagebox.showerror("Error", f"An error occurred while opening the file for update: {e}")

    def delete_file(self, note):
        file_path = note.get("path")
        if file_path:
            try:
                os.remove(file_path)
                self.populate_all_notes()  
            except Exception as e:
                messagebox.showerror("Error", f"An error occurred while deleting the file: {e}")

    def rename_file(self, note):
        file_path = note.get("path")
        if file_path:
            new_name = simpledialog.askstring("Rename", "Enter new name:")
            if new_name:
                new_path = os.path.join(os.path.dirname(file_path), new_name)
                try:
                    os.rename(file_path, new_path)
                    self.populate_all_notes() 
                except Exception as e:
                    messagebox.showerror("Error", f"An error occurred while renaming the file: {e}")

    def hide_all_notes_buttons(self):
        self.show_all_notes_button.grid_forget()

    ###############################
    # Section: Start page
    ###############################

    def hide_start_page(self):
        self.start_page.start_window.withdraw() 

    def initialize_main_app(self):
        self.hide_start_page()
        self.root.deiconify()  
        
if __name__ == "__main__":
    root = tk.Tk()
    app = NoteCraftApp(root, "vosk-model-small-en-in-0.4")
    icon_path = "images/pic.png"  
    root.iconbitmap(icon_path)
    root.mainloop()